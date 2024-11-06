import streamlit as st
from docx import Document
from io import BytesIO

# Global styling and page configuration
st.set_page_config(page_title="Startup Growth Metrics", layout="wide")
st.markdown("""
    <style>
    /* Custom CSS for the app */
    .navbar {
        display: flex;
        justify-content: space-around;
        background-color: #4CAF50;
        padding: 10px;
        color: white;
    }
    .navbar a {
        color: white;
        text-decoration: none;
        font-weight: bold;
        font-size: 1.2em;
    }
    .hero-section {
        text-align: center;
        padding: 60px 20px;
        color: white;
        background-color: #1E1E1E;
    }
    .hero-title {
        font-size: 3em;
        font-weight: bold;
    }
    .hero-subtitle {
        font-size: 1.2em;
        color: #A0A0A0;
    }
    .footer {
        text-align: center;
        padding: 10px;
        color: #A0A0A0;
        margin-top: 20px;
    }
    </style>
""", unsafe_allow_html=True)

# Navbar
page = st.sidebar.radio("Navigation", ["Home", "Metrics Guide", "Calculator"])

# Home Page
if page == "Home":
    # Hero Section
    st.markdown("<div class='hero-section'><h1 class='hero-title'>Startup Growth Metrics Calculator</h1><p class='hero-subtitle'>Calculate, analyze, and optimize your startup’s key metrics effortlessly.</p></div>", unsafe_allow_html=True)
    
    # Features Section
    st.subheader("🚀 Features")
    st.write("""
    - **Comprehensive Metrics**: Calculate essential metrics like CAC, LTV, ARPU, and more.
    - **Easy-to-Use Interface**: Intuitive and user-friendly design for hassle-free calculations.
    - **Downloadable Reports**: Generate and download a document with your metric results.
    """)
    
    # Footer
    st.markdown("<div class='footer'>Built to empower startups in optimizing financial performance.</div>", unsafe_allow_html=True)

# Metrics Guide Page
elif page == "Metrics Guide":
    st.title("📚 Metrics Guide")
    st.write("Learn more about each metric and understand their significance for your startup's growth.")

    # Documented descriptions for each metric
    metrics = {
        "Customer Acquisition Cost (CAC)": "Measures the cost of acquiring a new customer.",
        "Lifetime Value (LTV)": "Indicates the total revenue expected from a customer over their relationship with the business.",
        "Average Revenue Per User (ARPU)": "Represents the average revenue generated per user.",
        "LTV to CAC Ratio": "Helps assess the profitability of each customer.",
        "Break-Even Point (BEP)": "Indicates the number of units needed to cover both fixed and variable costs, achieving profitability.",
    }
    
    for metric, description in metrics.items():
        st.subheader(metric)
        st.write(description)

# Calculator Page
elif page == "Calculator":
    st.title("📐 Metrics Calculator")
    st.write("Select a metric to calculate and download the results.")

    # Metric Selection
    selected_metric = st.selectbox("Select a Metric", list(metrics.keys()))

    # Input fields based on selected metric
    result = None  # Placeholder for calculation result
    if selected_metric == "Customer Acquisition Cost (CAC)":
        marketing_expenses = st.number_input("Total Marketing Expenses", min_value=0.0)
        sales_expenses = st.number_input("Total Sales Expenses", min_value=0.0)
        customers_acquired = st.number_input("Number of Customers Acquired", min_value=1)
        if st.button("Calculate CAC"):
            result = (marketing_expenses + sales_expenses) / customers_acquired
            st.write(f"**CAC**: ${result:.2f}")
    
    elif selected_metric == "Lifetime Value (LTV)":
        purchase_value = st.number_input("Average Purchase Value", min_value=0.0)
        purchase_frequency = st.number_input("Average Purchase Frequency", min_value=0.0)
        customer_lifetime = st.number_input("Customer Lifetime", min_value=0.0)
        if st.button("Calculate LTV"):
            result = purchase_value * purchase_frequency * customer_lifetime
            st.write(f"**LTV**: ${result:.2f}")

    # Add similar blocks for other metrics...

    # Document Generation and Download
    if result is not None:
        # Create a downloadable document with the result
        def create_document(metric, value):
            doc = Document()
            doc.add_heading("Startup Growth Metrics Report", level=1)
            doc.add_paragraph(f"Metric: {metric}")
            doc.add_paragraph(f"Result: ${value:.2f}")
            doc.add_paragraph("Thank you for using the Startup Growth Metrics Calculator!")
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        buffer = create_document(selected_metric, result)
        st.download_button(
            label="📄 Download Report",
            data=buffer,
            file_name="metrics_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
